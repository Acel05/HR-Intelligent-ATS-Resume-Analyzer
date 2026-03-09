"""
Resume Parser Service - Extracts text and structured data from PDF/DOCX
"""
import re
from pypdf import PdfReader
from docx import Document
from typing import Dict, List, Any, Optional
from app.models.schemas import CandidateInfo, Project, Experience, ExperienceSummary, Education
from app.services.ocr_service import ocr_service


class ResumeParser:
    """Parse resumes and extract structured information"""
    
    PARSING_STANDARD = "standard"
    PARSING_OCR = "ocr"
    PARSING_OCR_UNAVAILABLE = "ocr_unavailable"
    
    # Regex patterns
    EMAIL_PATTERN = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    PHONE_PATTERN = r'(?:\+?1[-.\s]?)?(?:\(?\d{3}\)?[-.\s]?)?\d{3}[-.\s]?\d{4}|\+\d{1,3}[-.\s]?\d{6,14}'
    LINKEDIN_PATTERN = r'(?:linkedin\.com/in/|linkedin:?\s*)([a-zA-Z0-9-]+)'
    GITHUB_PATTERN = r'(?:github\.com/|github:?\s*)([a-zA-Z0-9-]+)'
    
    # Section headers (Kumpulan Variasi Ekstrem Bahasa Indonesia & Inggris)
    SECTION_HEADERS = {
        'experience': [
            'experience', 'work experience', 'employment', 'work history', 'professional experience', 
            'pengalaman', 'pengalaman kerja', 'riwayat pekerjaan', 'pengalaman profesional', 'riwayat kerja',
            'pengalaman organisasi', 'organisasi', 'kepanitiaan', 'pengalaman magang', 'internship'
        ],
        'education': [
            'education', 'academic', 'qualification', 'educational background', 
            'pendidikan', 'riwayat pendidikan', 'latar belakang pendidikan', 'edukasi', 
            'pendidikan formal', 'pendidikan non formal', 'academic background'
        ],
        'skills': [
            'skills', 'technical skills', 'competencies', 'technologies', 'expertise', 
            'keahlian', 'keterampilan', 'kemampuan', 'skill', 'kompetensi', 'kemampuan teknis',
            'bahasa', 'languages', 'hard skill', 'soft skill'
        ],
        'projects': [
            'projects', 'personal projects', 'key projects', 'portfolio', 
            'proyek', 'portofolio', 'project', 'karya', 'publikasi', 'pencapaian'
        ],
        'certifications': [
            'certifications', 'certificates', 'licenses', 
            'sertifikasi', 'lisensi', 'penghargaan', 'sertifikat', 'pelatihan', 'prestasi', 'awards', 'training'
        ],
        'summary': [
            'summary', 'profile', 'objective', 'about', 'professional summary', 
            'profil', 'tentang saya', 'ringkasan', 'objektif', 'ringkasan profesional', 'data pribadi', 'informasi pribadi'
        ]
    }
    
    # Action verbs (Ditambahkan puluhan kata kerja aktif Bahasa Indonesia)
    ACTION_VERBS = [
        'achieved', 'administered', 'analyzed', 'built', 'collaborated', 'created', 'delivered',
        'designed', 'developed', 'drove', 'enhanced', 'established', 'executed', 'implemented', 
        'improved', 'increased', 'integrated', 'launched', 'led', 'managed', 'mentored', 
        'optimized', 'orchestrated', 'oversaw', 'planned', 'reduced', 'resolved', 'spearheaded',
        'mencapai', 'mengelola', 'menganalisis', 'membangun', 'mengotomatisasi', 'berkolaborasi', 
        'membuat', 'mengirimkan', 'mendesain', 'merancang', 'mengembangkan', 'mendorong', 
        'meningkatkan', 'mendirikan', 'mengeksekusi', 'mengimplementasikan', 'menerapkan',
        'memimpin', 'membimbing', 'mengoptimalkan', 'merencanakan', 'mengurangi', 'menyelesaikan', 
        'mengawasi', 'menyelaraskan', 'menyusun', 'menangani', 'membantu', 'memastikan', 'melakukan',
        'bertanggung jawab', 'mengkoordinasi', 'mengatur'
    ]
    
    # Pola Tanggal Super Fleksibel (Bulan/Tahun Indo/Eng & Format Pemisah Indo)
    MONTHS_PATTERN = r'(?:Jan(?:uari)?|Feb(?:ruari)?|Mar(?:et)?|Apr(?:il)?|Mei|May|Jun(?:i)?|Jul(?:i)?|Agt|Agu(?:stus)?|Aug(?:ust)?|Sep(?:tember)?|Okt(?:ober)?|Oct(?:ober)?|Nov(?:ember)?|Des(?:ember)?|Dec(?:ember)?)'
    SEPARATOR_PATTERN = r'\s*(?:[-–—to]+|s/?d\.?|sampai(?: dengan)?|hingga)\s*'
    PRESENT_PATTERN = r'(?:Present|Current|Saat ini|Sekarang|Kini|Hingga kini)'
    
    def parse(self, file_path: str, file_ext: str) -> Dict[str, Any]:
        parsing_method = self.PARSING_STANDARD
        ocr_confidence = None
        
        if file_ext == '.pdf':
            raw_text = self._extract_pdf_text(file_path)
            has_tables = self._check_pdf_tables(file_path)
            has_images = self._check_pdf_images(file_path)
            
            raw_text, parsing_method, ocr_confidence = self._apply_ocr_if_needed(
                file_path, raw_text
            )
        else:
            raw_text = self._extract_docx_text(file_path)
            has_tables = self._check_docx_tables(file_path)
            has_images = self._check_docx_images(file_path)
        
        sections = self._identify_sections(raw_text)
        
        candidate = self._extract_candidate_info(raw_text)
        experience = self._extract_experience(raw_text, sections.get('experience', ''))
        projects = self._extract_projects(raw_text, sections.get('projects', ''))
        education = self._extract_education(raw_text, sections.get('education', ''))
        
        return {
            "raw_text": raw_text,
            "candidate": candidate,
            "experience": experience,
            "projects": projects,
            "education": education,
            "sections": sections,
            "formatting": {
                "has_tables": has_tables,
                "has_images": has_images,
                "word_count": len(raw_text.split()),
                "line_count": len(raw_text.split('\n'))
            },
            "parsing_method": parsing_method,
            "ocr_confidence": ocr_confidence
        }
    
    def _apply_ocr_if_needed(self, file_path: str, standard_text: str) -> tuple:
        if not ocr_service.is_available():
            return standard_text, self.PARSING_STANDARD, None
        
        email_match = re.search(self.EMAIL_PATTERN, standard_text)
        phone_match = re.search(self.PHONE_PATTERN, standard_text)
        
        if not ocr_service.needs_ocr(
            standard_text, 
            email=email_match.group() if email_match else None,
            phone=phone_match.group() if phone_match else None
        ):
            return standard_text, self.PARSING_STANDARD, None
        
        if ocr_service.should_skip_ocr(file_path):
            return standard_text, self.PARSING_OCR_UNAVAILABLE, None
        
        ocr_text, parsing_method, confidence = ocr_service.extract_text_with_ocr(file_path)
        
        if ocr_text and parsing_method == self.PARSING_OCR:
            return ocr_text, parsing_method, confidence
        else:
            return standard_text, parsing_method, confidence
    
    def _extract_pdf_text(self, file_path: str) -> str:
        text = ""
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")
        return text
    
    def _extract_docx_text(self, file_path: str) -> str:
        text = ""
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")
        return text
    
    def _check_pdf_tables(self, file_path: str) -> bool:
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                text = page.extract_text() or ""
                lines = text.split('\n')
                if sum(1 for line in lines if line.count('\t') >= 2 or line.count('|') >= 2) > 3:
                    return True
        except: pass
        return False
    
    def _check_pdf_images(self, file_path: str) -> bool:
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                if '/XObject' in page.get('/Resources', {}):
                    xobject = page['/Resources']['/XObject']
                    if xobject:
                        for obj in xobject:
                            if xobject[obj]['/Subtype'] == '/Image': return True
        except: pass
        return False
    
    def _check_docx_tables(self, file_path: str) -> bool:
        try:
            return len(Document(file_path).tables) > 0
        except: return False
    
    def _check_docx_images(self, file_path: str) -> bool:
        try:
            doc = Document(file_path)
            for rel in doc.part.rels.values():
                if "image" in rel.reltype: return True
        except: pass
        return False
    
    def _identify_sections(self, text: str) -> Dict[str, str]:
        sections = {}
        lines = text.split('\n')
        current_section = None
        current_content = []
        
        for line in lines:
            line_lower = line.lower().strip()
            section_found = None
            
            for section_type, headers in self.SECTION_HEADERS.items():
                for header in headers:
                    # Pencocokan lebih fleksibel agar terbaca meski ada karakter aneh/spasi berlebih
                    if line_lower == header or line_lower.startswith(header + ':') or line_lower.startswith(header + ' '):
                        section_found = section_type
                        break
                if section_found:
                    break
            
            if section_found:
                if current_section:
                    sections[current_section] = '\n'.join(current_content)
                current_section = section_found
                current_content = []
            elif current_section:
                current_content.append(line)
        
        if current_section:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
    
    def _extract_candidate_info(self, text: str) -> CandidateInfo:
        lines = text.split('\n')[:15]
        name = None
        
        for line in lines:
            line = line.strip()
            if len(line) > 2 and len(line) < 50:
                if '@' in line or re.search(self.PHONE_PATTERN, line):
                    continue
                if any(word in line.lower() for word in ['resume', 'cv', 'curriculum', 'profil', 'data pribadi', 'portofolio']):
                    continue
                words = line.split()
                if 1 <= len(words) <= 4 and all(w.replace('.', '').replace('-', '').replace(',', '').isalpha() for w in words):
                    name = line
                    break
        
        email_match = re.search(self.EMAIL_PATTERN, text)
        email = email_match.group() if email_match else None
        
        phone_match = re.search(self.PHONE_PATTERN, text)
        phone = phone_match.group() if phone_match else None
        
        linkedin_match = re.search(self.LINKEDIN_PATTERN, text, re.IGNORECASE)
        linkedin = f"linkedin.com/in/{linkedin_match.group(1)}" if linkedin_match else None
        
        github_match = re.search(self.GITHUB_PATTERN, text, re.IGNORECASE)
        github = f"github.com/{github_match.group(1)}" if github_match else None
        
        location = self._extract_location(text)
        
        return CandidateInfo(
            name=name, email=email, phone=phone, location=location, linkedin=linkedin, github=github
        )
    
    def _extract_location(self, text: str) -> Optional[str]:
        # Pola Lokasi Indonesia yang sangat umum
        location_patterns = [
            r'(?:Location|Address|Based in|City|Alamat|Domisili|Kota|Tempat Tinggal|Jl\.|Jalan)[:\s]*([A-Za-z0-9\s,\.]+)',
            r'([A-Za-z]+,\s*[A-Z]{2})\s*\d{5}',  
            r'([A-Za-z\s]+,\s*[A-Za-z\s]+,\s*(?:Indonesia))',  
        ]
        
        for pattern in location_patterns:
            match = re.search(pattern, text[:600], re.IGNORECASE)
            if match:
                location = match.group(1).strip()
                if len(location) > 3 and len(location) < 80:
                    return location
        return None
    
    def _extract_experience(self, full_text: str, experience_section: str) -> ExperienceSummary:
        positions = []
        text_to_analyze = experience_section if experience_section else full_text
        entries = self._split_experience_entries(text_to_analyze)
        
        total_months = 0
        
        for entry in entries[:5]: 
            exp = self._parse_experience_entry(entry)
            if exp.company or exp.role:
                positions.append(exp)
                if exp.duration:
                    months = self._estimate_duration_months(exp.duration)
                    total_months += months
        
        overall_quality = 0
        if positions:
            quality_sum = sum(p.bullet_quality for p in positions)
            overall_quality = quality_sum // len(positions)
        
        return ExperienceSummary(
            total_years=round(total_months / 12, 1),
            total_months=total_months,
            positions=positions,
            overall_quality=overall_quality
        )
    
    def _split_experience_entries(self, text: str) -> List[str]:
        # Mendeteksi format Bulan Tahun ATAU hanya Tahun saja (misal: 2020 - 2022)
        date_pattern = r'(' + self.MONTHS_PATTERN + r'\.?\s*\d{4}|\b\d{4}\b)'
        
        lines = text.split('\n')
        entries = []
        current_entry = []
        
        for line in lines:
            has_date = re.search(date_pattern, line, re.IGNORECASE)
            
            if has_date and current_entry:
                entries.append('\n'.join(current_entry))
                current_entry = []
            
            current_entry.append(line)
        
        if current_entry:
            entries.append('\n'.join(current_entry))
        
        return entries if entries else [text]
    
    def _parse_experience_entry(self, entry: str) -> Experience:
        lines = entry.strip().split('\n')
        
        company = None
        role = None
        duration = None
        
        for line in lines[:4]:
            line = line.strip()
            if not line:
                continue
            
            # Mendeteksi format durasi "Bulan Tahun - Bulan Tahun" atau "Tahun - Tahun"
            date_match = re.search(
                r'(' + self.MONTHS_PATTERN + r'\.?\s*\d{4}|\b\d{4}\b)' + self.SEPARATOR_PATTERN + r'(' + self.PRESENT_PATTERN + r'|' + self.MONTHS_PATTERN + r'\.?\s*\d{4}|\b\d{4}\b)',
                line, re.IGNORECASE
            )
            if date_match:
                duration = line
                continue
            
            # Jabatan Ekstrem Dwibahasa
            role_keywords = [
                'engineer', 'developer', 'manager', 'analyst', 'designer', 'lead', 'director', 'intern', 'associate', 'specialist', 'consultant',
                'staf', 'staff', 'manajer', 'analis', 'desainer', 'direktur', 'magang', 'spesialis', 'konsultan', 'asisten', 'teknisi', 
                'operator', 'admin', 'supervisor', 'ketua', 'wakil', 'anggota', 'koordinator', 'freelance', 'pekerja lepas', 'fasilitator'
            ]
            if any(kw in line.lower() for kw in role_keywords) and not role:
                role = line
            elif not company and len(line) > 2:
                # Perusahaan di Indonesia sering menggunakan PT, CV, dsb
                company = line
        
        bullets = [l for l in lines if l.strip().startswith(('•', '-', '*', '●', '○', '➢', '✓')) or (l.strip() and l.strip()[0].isdigit() and '.' in l[:3])]
        
        action_count = 0
        has_metrics = False
        for bullet in bullets:
            bullet_lower = bullet.lower()
            for verb in self.ACTION_VERBS:
                if verb in bullet_lower:
                    action_count += 1
                    break
            if re.search(r'\d+%|\$\d+|Rp\d+|increased|decreased|reduced|improved|meningkat|menurun|berkurang|sebanyak|mencapai', bullet_lower):
                has_metrics = True
        
        bullet_quality = 0
        if bullets:
            action_ratio = action_count / len(bullets)
            bullet_quality = int(action_ratio * 70)
            if has_metrics:
                bullet_quality += 30
            bullet_quality = min(100, bullet_quality)
        
        return Experience(
            company=company, role=role, duration=duration,
            description='\n'.join(bullets[:5]) if bullets else None,
            bullet_quality=bullet_quality, has_metrics=has_metrics, action_verbs_count=action_count
        )
    
    def _estimate_duration_months(self, duration_str: str) -> int:
        if re.search(self.PRESENT_PATTERN, duration_str, re.IGNORECASE):
            return 12
        
        # Ekstrak semua tahun yang ada di string
        dates = re.findall(r'\b(\d{4})\b', duration_str)
        
        if len(dates) >= 2:
            try:
                start_year = int(dates[0])
                end_year = int(dates[-1]) # Mengambil tahun terakhir
                return max(1, (end_year - start_year) * 12)
            except:
                pass
        
        return 12 
    
    def _extract_projects(self, full_text: str, projects_section: str) -> List[Project]:
        projects = []
        text_to_analyze = projects_section if projects_section else full_text
        entries = self._split_project_entries(text_to_analyze)
        
        for entry in entries[:5]:
            project = self._parse_project_entry(entry)
            if project.title:
                projects.append(project)
        
        return projects
    
    def _split_project_entries(self, text: str) -> List[str]:
        lines = text.split('\n')
        entries = []
        current_entry = []
        
        for line in lines:
            if line.strip() and not line.strip().startswith(('•', '-', '*', '●', '➢', '✓')):
                if current_entry and len(current_entry) > 1:
                    entries.append('\n'.join(current_entry))
                    current_entry = []
            current_entry.append(line)
        
        if current_entry:
            entries.append('\n'.join(current_entry))
        
        return entries
    
    def _parse_project_entry(self, entry: str) -> Project:
        lines = entry.strip().split('\n')
        title = None
        technologies = []
        description = []
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            if i == 0 or not title:
                title = line.replace('•', '').replace('-', '').strip()
            else:
                tech_match = re.search(r'(?:Tech|Technologies|Built with|Stack|Teknologi|Alat|Tools)[:\s]+(.+)', line, re.IGNORECASE)
                if tech_match:
                    techs = tech_match.group(1).split(',')
                    technologies.extend([t.strip() for t in techs])
                else:
                    description.append(line)
        
        score = 50
        if technologies:
            score += 20
        if description:
            score += 15
            desc_text = ' '.join(description).lower()
            if any(w in desc_text for w in ['improved', 'increased', 'reduced', 'users', 'revenue', 'meningkat', 'pengguna', 'berhasil', 'sukses', 'efisiensi']):
                score += 15
        
        return Project(
            title=title, technologies=technologies, description='\n'.join(description) if description else None,
            impact=None, score=min(100, score)
        )
    
    def _extract_education(self, full_text: str, education_section: str) -> List[Education]:
        education = []
        text_to_analyze = education_section if education_section else full_text
        
        lines = text_to_analyze.split('\n')
        current_edu = {}
        
        degree_keywords = [
            'bachelor', 'master', 'phd', 'doctorate', 'b.s.', 'b.a.', 'm.s.', 'm.a.', 'mba', 'b.tech', 'm.tech', 'b.e.', 'm.e.', 'diploma', 'associate',
            'sarjana', 'magister', 'doktor', 's1', 's2', 's3', 'd1', 'd2', 'd3', 'd4', 'sma', 'smk', 'stm', 'madya', 'gelar'
        ]
        
        # Seringkali nama institusi ditulis sebelum gelar
        institution_keywords = ['universitas', 'institut', 'politeknik', 'akademi', 'sekolah', 'university', 'college', 'institute', 'smk ', 'sma ']
        
        for line in lines:
            line_lower = line.lower().strip()
            if not line_lower:
                continue
            
            is_degree_line = False
            for keyword in degree_keywords:
                pattern = r'\b' + re.escape(keyword) + r'\b'
                if re.search(pattern, line_lower):
                    is_degree_line = True
                    if current_edu and ('degree' in current_edu or 'institution' in current_edu):
                        education.append(Education(**current_edu))
                        current_edu = {}
                    current_edu['degree'] = line.strip()
                    break
            
            if not is_degree_line:
                year_match = re.search(r'\b(19|20)\d{2}\b', line)
                if year_match and current_edu:
                    current_edu['year'] = year_match.group()
                
                # Menangani format IPK Indonesia yang sering menggunakan koma (misal: 3,85)
                gpa_match = re.search(r'(?:GPA|CGPA|IPK)[\s:]*(\d+[.,]\d+)', line, re.IGNORECASE)
                if gpa_match and current_edu:
                    current_edu['gpa'] = gpa_match.group(1).replace(',', '.')
                
                if current_edu and 'institution' not in current_edu and len(line.strip()) > 3:
                    current_edu['institution'] = line.strip()
                
                # Deteksi paksa untuk baris yang diawali nama Institusi Pendidikan
                elif not current_edu and any(inst in line_lower for inst in institution_keywords):
                    current_edu['institution'] = line.strip()
        
        if current_edu:
            education.append(Education(**current_edu))
        
        return education[:3]